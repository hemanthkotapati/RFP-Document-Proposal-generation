import os
import time
import google.generativeai as genai
from docx import Document
from dotenv import load_dotenv
import re

load_dotenv()
GEMINI_API_KEY =os.getenv("GEMINI_API_KEY_2")

# Configure Gemini API
genai.configure(api_key=GEMINI_API_KEY)

def upload_to_gemini(path, mime_type="application/pdf"):
    """Uploads a file to Gemini and returns the file object."""
    try:
        file = genai.upload_file(path, mime_type=mime_type)
        print(f"Uploaded file '{file.display_name}' as: {file.uri}")
        return file
    except Exception as e:
        print(f"Error uploading file {path}: {e}")
        return None

def wait_for_files_active(files):
    """Waits for the uploaded files to become active before use."""
    print("Waiting for file processing...")
    for name in (file.name for file in files):
        file = genai.get_file(name)
        while file.state.name == "PROCESSING":
            print(".", end="", flush=True)
            time.sleep(10)
            file = genai.get_file(name)
        if file.state.name != "ACTIVE":
            raise Exception(f"File {file.name} failed to process")
    print("...all files ready\n")

def extract_tables_and_text(text):

    # Split text at <br> (which separates tables in Gemini output)
    sections = re.split(r'<br\s*/?>', text)

    extracted_tables = []
    before_table = ""
    after_table = ""
    table_found = False

    for section in sections:
        lines = section.strip().split("\n")
        table_lines = []
        inside_table = False

        for line in lines:
            if "|" in line:  # Detect table start
                table_lines.append(line)
                inside_table = True
            elif inside_table and line.strip():  # Handle mis-formatted rows
                table_lines.append(line)
            elif inside_table and not line.strip():  # Table ends when empty line appears
                inside_table = False
            elif inside_table:
                table_lines.append(line)  # Handle any edge cases
            elif not table_found:  # Text before first table
                before_table += line + "\n"
            else:  # Text after last table
                after_table += line + "\n"

        if table_lines:
            extracted_tables.append("\n".join(table_lines).strip())
            table_found = True

    return before_table.strip(), extracted_tables, after_table.strip()


def add_table_to_word(doc, table_text):

    if not table_text:
        return doc  # No table detected

    lines = table_text.strip().split("\n")

    # Ensure at least two rows (headers + data)
    if len(lines) < 2:
        return doc

    # Extract headers and rows
    headers = lines[0].split("|")[1:-1]  # Remove first and last empty elements
    headers = [h.strip() for h in headers]

    rows = []
    for line in lines[2:]:  # Skip headers and separator row
        columns = line.split("|")[1:-1]
        if len(columns) == len(headers):  # Ensure row matches header count
            rows.append([col.strip() for col in columns])

    if not rows:
        return doc  # No valid table rows

    # Create Word table
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'

    # Add headers
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # Add rows
    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            table.cell(row_idx + 1, col_idx).text = cell_text

    return doc


# Define generation configuration
generation_config = {
    "temperature": 0.1,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

# Create the model
model = genai.GenerativeModel(
    model_name="gemini-2.0-flash",
    generation_config=generation_config,
    # system_instruction="Please respond to the given question from the content provided. Make sure to give a concise, structured and summarized response so that it aligns with the questions requirement.",
    # system_instruction="""You are an expert file search assistant. Please respond on the basis of the requirement of the
    #                       question from the content provided. Make sure to keep the response concise and structured.
    #                       Dont include example information or formulas in the response."""
    system_instruction = """
    You are an expert file search assistant. Provide concise and structured responses based solely on the provided content and the question's requirements. 
    Avoid including example information, formulas, or unnecessary details in your answers.
    """

)

#Don't include example information or formulas in the summary, but include tabular data wherever detected.
#If there are 2 or more tables in response the separate the tables with an empty line.

# Upload the document
files = [upload_to_gemini("../RFP_Documents/GAIL_Tender_Document.pdf", mime_type="application/pdf")]

# Wait for the files to be processed
wait_for_files_active(files)

# Start a chat session
chat_session = model.start_chat(history=[])

# Define multiple questions
# questions = [
#     # "How many acceptance tests do we need to do?",
#     # "What modules are required to be delivered?",
#     # "What functionalities are required in the software?",
#     # "How many servers are required to be supplied?",
#     # "How many interfaces are required?",
#     # "List out data of all the pipelines required to be configured?",
#     # "How many days of testing is required?",
#     # "How many training days are required?",
#     # "Give details of Customer Name, End User, Project Name.",
#     # "What is the requirement of this document?"
#     "What are all the modules to be provided and their description?"
# ]
#
# # Initialize Word Document
# doc = Document()
# doc.add_heading("Tender Document Summary", level=1)
#
# # Dictionary to store responses
# responses = {}
#
# # Iterate over questions and send them to the model
# for i, question in enumerate(questions):
#     print(f"Sending question {i+1}/{len(questions)}: {question}")
#
#     # Add question as a heading
#     doc.add_heading(question, level=2)
#
#     response = chat_session.send_message([files[0], question])
#     responses[question] = response.text
#
#     print(response.text)
#
#     # Extract text before, table, and text after
#     before_table_text, extracted_tables, after_table_text = extract_tables_and_text(response.text)
#
#     # Add text before table
#     if before_table_text:
#         doc.add_paragraph(before_table_text)
#
#     # Add table if detected
#     for table in extracted_tables:
#         doc = add_table_to_word(doc, table)
#         doc.add_paragraph("\n")
#
#     # Add text after table
#     if after_table_text:
#         doc.add_paragraph(after_table_text)
#
#     print(f"Response stored in document.\n")

questions_with_tags = [
    {"question": "What is the name of the customer?", "tag": "<<Customer Name>>"},
    {"question": "What is the project timeline?", "tag": "<<Project Timeline>>"},
    {"question": "What are the payment terms?", "tag": "<<Payment Terms>>"},
    {"question": "What are all the modules to be provided and their description?", "tag": "<<Modules and Description>>"}
]

# Initialize Word Document
doc = Document()
doc.add_heading("Tender Document Summary", level=1)

# Data list with question, tag, and response
data = []

# Iterate over questions, get responses, and store in data list
for item in questions_with_tags:
    question, tag = item["question"], item["tag"]
    print(f"Sending question: {question}")

    doc.add_heading(question, level=2)
    response = chat_session.send_message([files[0], question])

    entry = {"question": question, "tag": tag, "response": response.text}
    data.append(entry)

    before_text, tables, after_text = extract_tables_and_text(response.text)

    if before_text:
        doc.add_paragraph(before_text)
    for table in tables:
        add_table_to_word(doc, table)
        doc.add_paragraph("\n")
    if after_text:
        doc.add_paragraph(after_text)

    print(f"Response stored for tag {tag}.\n")

# Save the document
doc.save("../Generated_Docs/Tender_Document_Summary_table_011.docx")
print("Word document with full content saved successfully!")


# def extract_table_and_text(text):
#     """
#     Extracts markdown tables from mixed text and returns:
#     - Text before the table
#     - The extracted table
#     - Text after the table
#     """
#     lines = text.strip().split("\n")
#     before_table, table_lines, after_table = [], [], []
#     inside_table = False
#
#     for line in lines:
#         if "|" in line:  # Detect table start
#             table_lines.append(line)
#             inside_table = True
#         elif inside_table and line.strip():  # Handle mis-formatted rows
#             table_lines.append(line)
#         elif inside_table and not line.strip():  # Table ends when empty line appears
#             inside_table = False
#         elif inside_table:
#             table_lines.append(line)  # Handle any edge cases
#         elif not inside_table and not table_lines:  # Text before the table
#             before_table.append(line)
#         else:  # Text after the table
#             after_table.append(line)
#
#     return "\n".join(before_table).strip(), "\n".join(table_lines).strip(), "\n".join(after_table).strip()