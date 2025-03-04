import streamlit as st
import os
import time
import google.generativeai as genai
from dotenv import load_dotenv
import re
from docx import Document as WordDoc
from docx.oxml import OxmlElement
from docx2pdf import convert
import base64


# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY_2")

# Configure Gemini API
genai.configure(api_key=GEMINI_API_KEY)

# Upload file to Gemini
def upload_to_gemini(path, mime_type="application/pdf"):
    try:
        file = genai.upload_file(path, mime_type=mime_type)
        return file
    except Exception as e:
        st.error(f"Error uploading file: {e}")
        return None

# Wait for file processing
def wait_for_files_active(files):
    print("Waiting for file processing...")
    for name in (file.name for file in files):
        file = genai.get_file(name)
        while file.state.name == "PROCESSING":
            print(".", end="", flush=True)
            time.sleep(10)
            file = genai.get_file(name)
        if file.state.name != "ACTIVE":
            raise Exception(f"File {file.name} failed to process")
    print("...all files ready\n")

def convert_docx_to_pdf(docx_path):
    """ Convert .docx to .pdf using docx2pdf """
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    return pdf_path

def save_uploaded_file(uploaded_file):
    """ Save uploaded file and convert .docx to .pdf if needed """
    temp_path = os.path.join("Uploaded_Docs", uploaded_file.name)
    os.makedirs("Uploaded_Docs", exist_ok=True)

    with open(temp_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    if uploaded_file.name.endswith(".docx"):
        pdf_path = convert_docx_to_pdf(temp_path)
        return pdf_path

    return temp_path

# def save_uploaded_file(uploaded_file):
#     temp_path = os.path.join("Uploaded_Docs", uploaded_file.name)
#     os.makedirs("Uploaded_Docs", exist_ok=True)
#
#     with open(temp_path, "wb") as f:
#         f.write(uploaded_file.getbuffer())
#
#     return temp_path


# Streamlit UI
st.title("AI RFP Processing with Gemini")

# File uploader
uploaded_file = st.file_uploader("Upload your RFP document", type=["pdf", "docx"])

if uploaded_file:
    st.success(f"File uploaded: {uploaded_file.name}")

    file_path = save_uploaded_file(uploaded_file)

    gemini_file = upload_to_gemini(file_path)

    if gemini_file:
        st.success("File uploaded and processing started...")
        wait_for_files_active([gemini_file])

        # Configure Gemini model
        generation_config = {
            "temperature": 0.1,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }

        model = genai.GenerativeModel(
            model_name="gemini-2.0-flash",
            generation_config=generation_config,
            system_instruction="""
            You are an expert file search assistant. Provide concise and structured responses based solely on the provided content and the query's requirements. 
            Do not ignore related content to the query as the information is very critical. If asked to list out details about anything make sure to include all the data related to it.
            Avoid including example information, formulas, or unnecessary details in your answers. Make sure to include tabular information wherever detected.
            """
        )

        chat_session = model.start_chat(history=[])

        # Questions with tags
        questions_with_tags = [
            {"question": "List out all the application software modules to be provided.", "tag": "<<Modules>>"},
            {"question": "List out all details of all the pipelines required to be configured in a table.", "tag": "<<Scope of Assets>>"},
            {"question": "List out all the deliverables from the APPS Vendor side.", "tag": "<<Deliverables>>"},
            {"question": "What are all the works to be performed for the customer assets?", "tag": "<<Work to be performed>>"},
        ]

        # Process questions
        responses = []

        for i, item in enumerate(questions_with_tags):
            question, tag = item["question"], item["tag"]
            print(f"Sending question {i + 1}/{len(questions_with_tags)}: {question}")

            response = chat_session.send_message([gemini_file, question])
            response_text = response.text.strip()
            responses.append({"question": question, "tag": tag, "response": response_text})

            print("Response: ", response_text)


        # Extract tables and surrounding text
        def extract_tables_and_text(text):
            sections = re.split(r'<br\s*/?>', text)
            extracted_tables, before_table, after_table = [], "", ""
            table_found = False

            for section in sections:
                lines = section.strip().split("\n")
                table_lines, inside_table = [], False

                for line in lines:
                    if "|" in line:
                        table_lines.append(line)
                        inside_table = True
                    elif inside_table and line.strip():
                        table_lines.append(line)
                    elif inside_table and not line.strip():
                        inside_table = False
                    elif not table_found:
                        before_table += line + "\n"
                    else:
                        after_table += line + "\n"

                if table_lines:
                    extracted_tables.append("\n".join(table_lines).strip())
                    table_found = True

            return before_table.strip(), extracted_tables, after_table.strip()


        # Insert an element after a paragraph (helper function)
        def insert_element_after(paragraph, element):
            p = paragraph._p
            p.addnext(element)


        # Add table after a paragraph
        def add_table_after_paragraph(paragraph, table_text):
            lines = table_text.strip().split("\n")
            if len(lines) < 2:
                return

            headers = [h.strip() for h in lines[0].split("|")[1:-1]]
            rows = [
                [col.strip() for col in line.split("|")[1:-1]]
                for line in lines[2:]
                if len(line.split("|")) - 2 == len(headers)
            ]

            if not rows:
                return

            # Create a temporary document to build the table
            temp_doc = WordDoc()
            table = temp_doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = 'Table Grid'

            # Fill headers
            for i, header in enumerate(headers):
                table.cell(0, i).text = header

            # Fill rows
            for row_idx, row in enumerate(rows):
                for col_idx, cell_text in enumerate(row):
                    table.cell(row_idx + 1, col_idx).text = cell_text

            # Insert the table after the target paragraph
            insert_element_after(paragraph, table._element)


        # Add a new paragraph after a paragraph
        def add_paragraph_after(paragraph, text):
            new_paragraph = OxmlElement('w:p')
            insert_element_after(paragraph, new_paragraph)
            p = paragraph._parent.add_paragraph(text)
            insert_element_after(paragraph, p._element)
            return p


        # Fill template with responses
        def fill_template(template_path, output_path, response_data):
            doc = WordDoc(template_path)

            for item in response_data:
                tag, response = item["tag"], item["response"].strip()
                before_text, tables, after_text = extract_tables_and_text(response)

                for para in doc.paragraphs:
                    if tag in para.text:
                        # Remove the tag from the paragraph
                        para.text = para.text.replace(tag, "")

                        # Add before_text to the same paragraph
                        if before_text:
                            para.add_run(before_text.strip())

                        # Add tables after the paragraph
                        for table_text in tables:
                            add_table_after_paragraph(para, table_text)

                        # Add after_text as a new paragraph
                        if after_text:
                            add_paragraph_after(para, after_text.strip())

            doc.save(output_path)
            print(f"Template filled and saved to {output_path}")

        # Paths
        template_path = "../Proposal_Documents/Emerson_Proposal_Template.docx"
        output_path = "../Proposal_Documents/streamlit_filled_template_4.docx"

        # Process Template
        fill_template(template_path, output_path, responses)
        st.success("Template filled successfully!")

        # Convert .docx to PDF
        output_pdf_path = output_path.replace(".docx", ".pdf")
        convert(output_path, output_pdf_path)

        # Display download button for PDF
        with open(output_path, "rb") as file:
            btn = st.download_button(
                label="Download Proposal Document",
                data=file,
                file_name="Processed_RFP_Template.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # Display PDF in Streamlit
        with open(output_pdf_path, "rb") as pdf_file:
            base64_pdf = base64.b64encode(pdf_file.read()).decode('utf-8')
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="800" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
